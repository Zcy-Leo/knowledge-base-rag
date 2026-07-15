import os
import json
import csv
import numpy as np
from scipy.stats import ttest_rel

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(BASE_DIR, "experiment_results")

SEARCH_MODES = [
    {"name": "Pure_Vector", "mode": "vector", "reranker": None, "model": None},
    {"name": "Pure_BM25", "mode": "bm25", "reranker": None, "model": None},
    {"name": "Hybrid_RRF", "mode": "hybrid", "reranker": None, "model": None},
    {"name": "Hybrid_CrossEncoder_L6", "mode": "hybrid", "reranker": "crossencoder", "model": "cross-encoder/ms-marco-MiniLM-L-6-v2"},
    {"name": "Hybrid_CrossEncoder_L12", "mode": "hybrid", "reranker": "crossencoder", "model": "cross-encoder/ms-marco-MiniLM-L-12-v2"},
    {"name": "Hybrid_LLM_DeepSeek", "mode": "hybrid", "reranker": "llm", "model": "deepseek"},
]

def cohens_d(group1, group2):
    diff = np.mean(group1) - np.mean(group2)
    pooled_std = np.sqrt((np.std(group1, ddof=1)**2 + np.std(group2, ddof=1)**2) / 2)
    return diff / pooled_std if pooled_std > 0 else 0

def load_metrics():
    metrics_path = os.path.join(OUTPUT_DIR, "experiment_metrics_v4.csv")
    all_metrics = []
    with open(metrics_path, 'r', encoding='utf-8') as f:
        reader = csv.DictReader(f)
        for row in reader:
            row['latency'] = float(row['latency'])
            row['ndcg@5'] = float(row['ndcg@5'])
            row['ndcg@10'] = float(row['ndcg@10'])
            row['mrr'] = float(row['mrr'])
            row['recall@5'] = float(row['recall@5'])
            row['recall@10'] = float(row['recall@10'])
            row['precision@5'] = float(row['precision@5'])
            row['precision@10'] = float(row['precision@10'])
            row['f1@5'] = float(row['f1@5'])
            row['f1@10'] = float(row['f1@10'])
            row['relevant_count'] = int(float(row['relevant_count']))
            all_metrics.append(row)
    return all_metrics

def compute_summary(all_metrics):
    summary = {}
    for mode_config in SEARCH_MODES:
        mode_name = mode_config['name']
        mode_metrics = [m for m in all_metrics if m['mode'] == mode_name]
        
        avg_latency = np.mean([m['latency'] for m in mode_metrics])
        avg_ndcg5 = np.mean([m['ndcg@5'] for m in mode_metrics])
        avg_ndcg10 = np.mean([m['ndcg@10'] for m in mode_metrics])
        avg_mrr = np.mean([m['mrr'] for m in mode_metrics])
        avg_recall5 = np.mean([m['recall@5'] for m in mode_metrics])
        avg_recall10 = np.mean([m['recall@10'] for m in mode_metrics])
        avg_precision5 = np.mean([m['precision@5'] for m in mode_metrics])
        avg_precision10 = np.mean([m['precision@10'] for m in mode_metrics])
        avg_f15 = np.mean([m['f1@5'] for m in mode_metrics])
        avg_f110 = np.mean([m['f1@10'] for m in mode_metrics])
        
        summary[mode_name] = {
            'avg_latency': avg_latency,
            'avg_ndcg@5': avg_ndcg5,
            'avg_ndcg@10': avg_ndcg10,
            'avg_mrr': avg_mrr,
            'avg_recall@5': avg_recall5,
            'avg_recall@10': avg_recall10,
            'avg_precision@5': avg_precision5,
            'avg_precision@10': avg_precision10,
            'avg_f1@5': avg_f15,
            'avg_f1@10': avg_f110,
            'total_queries': len(mode_metrics),
            'std_ndcg@10': np.std([m['ndcg@10'] for m in mode_metrics]),
            'std_mrr': np.std([m['mrr'] for m in mode_metrics]),
            'std_latency': np.std([m['latency'] for m in mode_metrics]),
        }
    return summary

def compute_statistical_significance(all_metrics):
    mode_names = [m['name'] for m in SEARCH_MODES]
    pairwise_results = []
    
    for i, mode1 in enumerate(mode_names):
        for j, mode2 in enumerate(mode_names):
            if i >= j:
                continue
            
            m1_metrics = [m for m in all_metrics if m['mode'] == mode1]
            m2_metrics = [m for m in all_metrics if m['mode'] == mode2]
            
            m1_ndcg = [m['ndcg@10'] for m in m1_metrics]
            m2_ndcg = [m['ndcg@10'] for m in m2_metrics]
            
            m1_mrr = [m['mrr'] for m in m1_metrics]
            m2_mrr = [m['mrr'] for m in m2_metrics]
            
            t_stat_ndcg, p_val_ndcg = ttest_rel(m1_ndcg, m2_ndcg)
            t_stat_mrr, p_val_mrr = ttest_rel(m1_mrr, m2_mrr)
            
            d_ndcg = cohens_d(m1_ndcg, m2_ndcg)
            d_mrr = cohens_d(m1_mrr, m2_mrr)
            
            effect_size_ndcg = "small" if abs(d_ndcg) < 0.2 else "medium" if abs(d_ndcg) < 0.5 else "large"
            effect_size_mrr = "small" if abs(d_mrr) < 0.2 else "medium" if abs(d_mrr) < 0.5 else "large"
            
            pairwise_results.append({
                'mode1': mode1,
                'mode2': mode2,
                't_stat_ndcg': float(t_stat_ndcg),
                'p_val_ndcg': float(p_val_ndcg),
                'effect_size_ndcg': effect_size_ndcg,
                'cohens_d_ndcg': float(d_ndcg),
                't_stat_mrr': float(t_stat_mrr),
                'p_val_mrr': float(p_val_mrr),
                'effect_size_mrr': effect_size_mrr,
                'cohens_d_mrr': float(d_mrr),
                'significant_ndcg': bool(p_val_ndcg < 0.05),
                'significant_mrr': bool(p_val_mrr < 0.05),
            })
            
            print(f"  {mode1} vs {mode2}:")
            print(f"    NDCG@10: t={t_stat_ndcg:.3f}, p={p_val_ndcg:.4f}, d={d_ndcg:.3f} ({effect_size_ndcg})")
            print(f"    MRR:     t={t_stat_mrr:.3f}, p={p_val_mrr:.4f}, d={d_mrr:.3f} ({effect_size_mrr})")
    
    return pairwise_results

def generate_report(summary, pairwise_results):
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    title = doc.add_heading('Experimental Evaluation of Search and Reranking Models', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph('This paper presents a comprehensive experimental evaluation of different search and reranking models within a knowledge base automation system. We compare six search configurations: Pure Vector Search, Pure BM25, Hybrid RRF, Hybrid + CrossEncoder (MiniLM-L6), Hybrid + CrossEncoder (MiniLM-L12), and Hybrid + LLM (DeepSeek). The evaluation is conducted on a dataset of 2483 documents across 50 test queries, measuring performance using NDCG, MRR, Recall, Precision, and F1 metrics.')
    
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph('With the exponential growth of enterprise knowledge bases, effective information retrieval systems have become crucial. Traditional keyword-based search (BM25) and modern semantic search (Vector) each have their limitations. Hybrid approaches combining both methods with reranking have emerged as promising solutions. This work evaluates multiple retrieval strategies to identify optimal configurations for enterprise knowledge retrieval.')
    
    doc.add_heading('2. Methodology', level=1)
    
    doc.add_heading('2.1 Dataset', level=2)
    doc.add_paragraph('The experimental dataset consists of 2483 documents extracted from various enterprise PDF manuals and guides. Documents are preprocessed into text chunks with corresponding embeddings stored in FAISS index.')
    
    doc.add_heading('2.2 Evaluation Queries', level=2)
    doc.add_paragraph('Fifty test queries are designed, covering easy, medium, and hard difficulty levels. Each query has manually annotated relevant documents for ground truth evaluation.')
    
    doc.add_heading('2.3 Evaluation Metrics', level=2)
    doc.add_paragraph('• NDCG@5 and NDCG@10: Normalized Discounted Cumulative Gain at ranks 5 and 10')
    doc.add_paragraph('• MRR: Mean Reciprocal Rank')
    doc.add_paragraph('• Recall@5 and Recall@10: Proportion of relevant documents retrieved')
    doc.add_paragraph('• Precision@5 and Precision@10: Proportion of retrieved documents that are relevant')
    doc.add_paragraph('• F1@5 and F1@10: Harmonic mean of precision and recall')
    doc.add_paragraph('• Latency: Query processing time in seconds')
    
    doc.add_heading('2.4 Search Models', level=2)
    doc.add_paragraph('Six search configurations are evaluated:')
    doc.add_paragraph('1. Pure_Vector: Vector similarity search using FAISS')
    doc.add_paragraph('2. Pure_BM25: BM25 keyword-based search')
    doc.add_paragraph('3. Hybrid_RRF: Hybrid search using RRF fusion')
    doc.add_paragraph('4. Hybrid_CrossEncoder_L6: RRF + MiniLM-L-6-v2 reranker')
    doc.add_paragraph('5. Hybrid_CrossEncoder_L12: RRF + MiniLM-L-12-v2 reranker')
    doc.add_paragraph('6. Hybrid_LLM_DeepSeek: RRF + DeepSeek LLM reranker')
    
    doc.add_heading('3. Results', level=1)
    
    doc.add_heading('3.1 Overall Performance Comparison', level=2)
    
    headers1 = ['Search Mode', 'Avg. Latency (s)', 'NDCG@5', 'NDCG@10', 'MRR', 'Recall@10', 'Precision@10', 'F1@10']
    data1 = []
    
    configs = [
        ('Pure Vector', 'Pure_Vector'),
        ('BM25', 'Pure_BM25'),
        ('Hybrid (RRF)', 'Hybrid_RRF'),
        ('Hybrid + MiniLM-L6', 'Hybrid_CrossEncoder_L6'),
        ('Hybrid + MiniLM-L12', 'Hybrid_CrossEncoder_L12'),
        ('Hybrid + LLM (DeepSeek)', 'Hybrid_LLM_DeepSeek'),
    ]
    
    for name, key in configs:
        s = summary[key]
        data1.append([
            name,
            f"{s['avg_latency']:.3f}",
            f"{s['avg_ndcg@5']:.4f}",
            f"{s['avg_ndcg@10']:.4f}",
            f"{s['avg_mrr']:.4f}",
            f"{s['avg_recall@10']:.3f}",
            f"{s['avg_precision@10']:.3f}",
            f"{s['avg_f1@10']:.3f}",
        ])
    
    table1 = doc.add_table(rows=1, cols=len(headers1))
    table1.style = 'Table Grid'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = table1.rows[0].cells
    for i, header in enumerate(headers1):
        hdr_cells[i].text = header
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    for row_data in data1:
        row_cells = table1.add_row().cells
        for i, data in enumerate(row_data):
            row_cells[i].text = data
            row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    doc.add_paragraph()
    
    doc.add_heading('3.2 Statistical Significance Analysis', level=2)
    doc.add_paragraph('Pairwise comparisons using paired t-tests and Cohen\'s d effect size:')
    
    headers2 = ['Comparison', 'NDCG@10 t-stat', 'NDCG@10 p-value', 'NDCG@10 Effect Size', 'MRR t-stat', 'MRR p-value', 'MRR Effect Size']
    data2 = []
    
    for result in pairwise_results:
        sig_ndcg = "*" if result['significant_ndcg'] else ""
        sig_mrr = "*" if result['significant_mrr'] else ""
        data2.append([
            f"{result['mode1']} vs {result['mode2']}",
            f"{result['t_stat_ndcg']:.3f}",
            f"{result['p_val_ndcg']:.4f}{sig_ndcg}",
            f"{result['effect_size_ndcg']}",
            f"{result['t_stat_mrr']:.3f}",
            f"{result['p_val_mrr']:.4f}{sig_mrr}",
            f"{result['effect_size_mrr']}",
        ])
    
    table2 = doc.add_table(rows=1, cols=len(headers2))
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = table2.rows[0].cells
    for i, header in enumerate(headers2):
        hdr_cells[i].text = header
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    for row_data in data2:
        row_cells = table2.add_row().cells
        for i, data in enumerate(row_data):
            row_cells[i].text = data
            row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    doc.add_paragraph('* p < 0.05, statistically significant')
    
    doc.add_paragraph()
    
    doc.add_heading('3.3 Performance Summary', level=2)
    
    best_ndcg = max(summary.items(), key=lambda x: x[1]['avg_ndcg@10'])
    best_mrr = max(summary.items(), key=lambda x: x[1]['avg_mrr'])
    best_speed = min(summary.items(), key=lambda x: x[1]['avg_latency'])
    best_f1 = max(summary.items(), key=lambda x: x[1]['avg_f1@10'])
    
    doc.add_paragraph(f"• Best NDCG@10: {best_ndcg[0]} ({best_ndcg[1]['avg_ndcg@10']:.4f})")
    doc.add_paragraph(f"• Best MRR: {best_mrr[0]} ({best_mrr[1]['avg_mrr']:.4f})")
    doc.add_paragraph(f"• Fastest: {best_speed[0]} ({best_speed[1]['avg_latency']:.3f}s)")
    doc.add_paragraph(f"• Best F1@10: {best_f1[0]} ({best_f1[1]['avg_f1@10']:.3f})")
    
    doc.add_paragraph()
    
    doc.add_heading('4. Discussion', level=1)
    doc.add_paragraph('The experimental results show that BM25 outperforms pure vector search in terms of NDCG@10 (0.218 vs 0.188) and F1@10 (0.220 vs 0.159), with statistically significant differences observed (p=0.015). This suggests that keyword matching remains effective for structured enterprise documentation.')
    doc.add_paragraph('The CrossEncoder rerankers (L6 and L12) did not improve over baseline methods, likely due to the relatively small candidate set size and domain mismatch between the general-domain MS MARCO trained models and enterprise documentation.')
    doc.add_paragraph('The LLM-based reranker (DeepSeek) achieved competitive MRR (0.383) comparable to vector search and hybrid RRF, but with higher latency (12.62s per query) due to API call overhead.')
    doc.add_paragraph('For practical deployment, Hybrid RRF provides the best balance between retrieval effectiveness and efficiency, with average latency of only 0.035s while maintaining competitive MRR of 0.385.')
    
    doc.add_heading('5. Conclusion', level=1)
    doc.add_paragraph('This study systematically evaluates six search configurations for enterprise knowledge retrieval. The findings indicate that BM25 keyword search remains a strong baseline, while hybrid approaches combining vector and keyword search offer robustness. LLM-based reranking shows potential but requires optimization for production deployment due to latency concerns.')
    
    doc.add_paragraph()
    
    footer = doc.sections[0].footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = 'Experimental Evaluation Report - Knowledge Base Automation System'
    
    report_path = os.path.join(OUTPUT_DIR, "Experimental_Evaluation_Paper_v4.docx")
    doc.save(report_path)
    print(f"\n✅ Report saved to {report_path}")

def main():
    print("=== Loading Saved Metrics ===")
    all_metrics = load_metrics()
    print(f"Loaded {len(all_metrics)} metric records")
    
    print("\n=== Computing Summary ===")
    summary = compute_summary(all_metrics)
    
    print("\n=== Computing Statistical Significance ===")
    pairwise_results = compute_statistical_significance(all_metrics)
    
    print("\n=== Saving Data Files ===")
    summary_path = os.path.join(OUTPUT_DIR, "experiment_summary_v4.json")
    with open(summary_path, 'w', encoding='utf-8') as f:
        json.dump(summary, f, indent=2)
    print(f"✅ Summary saved to {summary_path}")
    
    significance_path = os.path.join(OUTPUT_DIR, "statistical_significance_v4.json")
    with open(significance_path, 'w', encoding='utf-8') as f:
        json.dump(pairwise_results, f, indent=2)
    print(f"✅ Statistical significance saved to {significance_path}")
    
    print("\n=== Generating Report ===")
    generate_report(summary, pairwise_results)
    
    print("\n" + "="*50)
    print("REPORT GENERATION COMPLETE")
    print("="*50)

if __name__ == "__main__":
    main()
