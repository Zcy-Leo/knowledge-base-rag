import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def set_cell_color(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    for edge in ['top', 'left', 'bottom', 'right']:
        element = parse_xml(f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="4" w:space="0" w:color="000000"/>')
        tcPr.append(element)

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

title = doc.add_heading('Weekly Progress Report', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_format = title.paragraph_format
title_format.space_after = Pt(12)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = subtitle.add_run('Knowledge Base Automation System')
sr.font.size = Pt(14)
sr.bold = True
sr.font.color.rgb = RGBColor(0, 51, 102)

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr2 = subtitle2.add_run('Period: July 8 - July 14, 2026')
sr2.font.size = Pt(11)
sr2.font.color.rgb = RGBColor(102, 102, 102)

doc.add_paragraph()

doc.add_heading('1. Executive Summary', level=1)
p1 = doc.add_paragraph()
p1.paragraph_format.first_line_indent = Cm(0.74)
p1.add_run('This weekly report summarizes the progress made on two key tasks: ')
p1.add_run('\n\n')
p1.add_run('1. ')
p1.add_run('Evaluate different search and reranking models and analyze their performance.').bold = True
p1.add_run(' Experiments were conducted on 6 search modes using 2483 documents and 50 test queries, generating statistically validated performance metrics.')
p1.add_run('\n\n')
p1.add_run('2. ')
p1.add_run('Review the LiveVectorLake framework and enhance the retrieval mechanism within the system.').bold = True
p1.add_run(' Significant enhancements were made to the incremental update mechanism, FAISS indexing, and search performance optimization.')
p1.add_run('\n\n')
p1.add_run('Due to network connectivity issues in China, the Gemini API was replaced with the DeepSeek API for LLM-based reranking experiments.')

doc.add_heading('2. Search and Reranking Model Evaluation', level=1)

doc.add_heading('2.1 Experimental Setup', level=2)
p2 = doc.add_paragraph()
p2.paragraph_format.first_line_indent = Cm(0.74)
p2.add_run('• **Dataset:** 2483 documents from the knowledge base')
p2.add_run('\n• **Test Queries:** 50 manually constructed queries covering various knowledge domains')
p2.add_run('\n• **Evaluation Metrics:** NDCG@5, NDCG@10, MRR, Recall@5, Recall@10, Precision@5, Precision@10, F1@5, F1@10, Latency')
p2.add_run('\n• **Statistical Methods:** Paired t-test, Cohen\'s d effect size analysis')

doc.add_heading('2.2 Search Modes Evaluated', level=2)
table1 = doc.add_table(rows=7, cols=4)
table1.alignment = WD_TABLE_ALIGNMENT.CENTER
table1.autofit = True

h = table1.rows[0].cells
headers = ['Search Mode', 'Mode', 'Reranker', 'Model']
for i, header in enumerate(headers):
    h[i].text = header
    h[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_color(h[i], '4472C4')
    for paragraph in h[i].paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True

data1 = [
    ['Pure_Vector', 'Vector', '-', '-'],
    ['Pure_BM25', 'BM25', '-', '-'],
    ['Hybrid_RRF', 'Hybrid', 'RRF', '-'],
    ['Hybrid_CrossEncoder_L6', 'Hybrid', 'CrossEncoder', 'ms-marco-MiniLM-L-6-v2'],
    ['Hybrid_CrossEncoder_L12', 'Hybrid', 'CrossEncoder', 'ms-marco-MiniLM-L-12-v2'],
    ['Hybrid_LLM_DeepSeek', 'Hybrid', 'LLM', 'DeepSeek-chat'],
]

for i, row in enumerate(data1):
    cells = table1.rows[i+1].cells
    for j, val in enumerate(row):
        cells[j].text = val
        cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if i % 2 == 0:
            set_cell_color(cells[j], 'D9E2F3')

doc.add_heading('2.3 Performance Results', level=2)
table2 = doc.add_table(rows=7, cols=8)
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
table2.autofit = True

h2 = table2.rows[0].cells
headers2 = ['Search Mode', 'Latency (s)', 'NDCG@5', 'NDCG@10', 'MRR', 'Recall@10', 'Precision@10', 'F1@10']
for i, header in enumerate(headers2):
    h2[i].text = header
    h2[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_color(h2[i], '4472C4')
    for paragraph in h2[i].paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True

data2 = [
    ['Pure_Vector', '0.0266', '0.2318', '0.1884', '0.3841', '0.3400', '0.1040', '0.1591'],
    ['Pure_BM25', '0.0090', '0.2565', '0.2180', '0.3476', '0.4667', '0.1440', '0.2198'],
    ['Hybrid_RRF', '0.0346', '0.2318', '0.1890', '0.3845', '0.3467', '0.1060', '0.1622'],
    ['Hybrid_CrossEncoder_L6', '5.5715', '0.1972', '0.1615', '0.3421', '0.2983', '0.0920', '0.1404'],
    ['Hybrid_CrossEncoder_L12', '11.4134', '0.1971', '0.1624', '0.3558', '0.2850', '0.0880', '0.1343'],
    ['Hybrid_LLM_DeepSeek', '12.6236', '0.2220', '0.1822', '0.3832', '0.3333', '0.1020', '0.1560'],
]

for i, row in enumerate(data2):
    cells = table2.rows[i+1].cells
    for j, val in enumerate(row):
        cells[j].text = val
        cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if i % 2 == 0:
            set_cell_color(cells[j], 'D9E2F3')

doc.add_heading('2.4 Key Findings', level=2)
p3 = doc.add_paragraph()
p3.paragraph_format.first_line_indent = Cm(0.74)
p3.add_run('• **Best Overall Performance (NDCG@10):** Pure_BM25 (0.2180) - Significantly outperforms CrossEncoder-based reranking')
p3.add_run('\n• **Best MRR:** Hybrid_RRF (0.3845) and Pure_Vector (0.3841) - Tied for first place')
p3.add_run('\n• **Best Recall:** Pure_BM25 (0.4667) - Almost 40% higher than CrossEncoder_L12')
p3.add_run('\n• **Fastest:** Pure_BM25 (0.009s) - 3x faster than Pure_Vector')
p3.add_run('\n• **Slowest:** Hybrid_LLM_DeepSeek (12.62s) - Due to API call overhead')
p3.add_run('\n• **CrossEncoder Performance:** L6 and L12 models show similar performance, with L6 being 2x faster')
p3.add_run('\n• **Hybrid RRF:** Performance is nearly identical to Pure_Vector, indicating limited benefit from RRF fusion')

doc.add_heading('2.5 Statistical Significance Analysis', level=2)
p4 = doc.add_paragraph()
p4.add_run('Significant Differences (p < 0.05):')

table3 = doc.add_table(rows=4, cols=5)
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
table3.autofit = True

h3 = table3.rows[0].cells
headers3 = ['Comparison', 'Metric', 't-statistic', 'p-value', 'Cohen\'s d']
for i, header in enumerate(headers3):
    h3[i].text = header
    h3[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_color(h3[i], '4472C4')
    for paragraph in h3[i].paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True

data3 = [
    ['Pure_BM25 vs CrossEncoder_L6', 'NDCG@10', '2.524', '0.0149*', '0.217 (medium)'],
    ['Pure_BM25 vs CrossEncoder_L12', 'NDCG@10', '2.319', '0.0246*', '0.211 (medium)'],
    ['Hybrid_RRF vs CrossEncoder_L6', 'NDCG@10', '2.047', '0.0460*', '0.118 (small)'],
]

for i, row in enumerate(data3):
    cells = table3.rows[i+1].cells
    for j, val in enumerate(row):
        cells[j].text = val
        cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if i % 2 == 0:
            set_cell_color(cells[j], 'D9E2F3')

p5 = doc.add_paragraph()
p5r = p5.add_run('* Statistically significant at p < 0.05 level')
p5r.font.size = Pt(10)
p5r.font.color.rgb = RGBColor(102, 102, 102)

doc.add_heading('3. LiveVectorLake Framework Enhancement', level=1)

doc.add_heading('3.1 Incremental Update Mechanism', level=2)
p6 = doc.add_paragraph()
p6.paragraph_format.first_line_indent = Cm(0.74)
p6.add_run('**Architecture:** Refactored from block-level to entry-level change detection')
p6.add_run('\n\n**Key Features:**')
p6.add_run('\n• SHA-256 content hash for each entry')
p6.add_run('\n• Detection of new, modified, and deleted entries')
p6.add_run('\n• Version number tracking and history management')
p6.add_run('\n• Efficient delta processing for incremental ingestion')

doc.add_heading('3.2 FAISS Index Optimization', level=2)
p7 = doc.add_paragraph()
p7.paragraph_format.first_line_indent = Cm(0.74)
p7.add_run('• **Cache Mechanism:** Implemented st.cache_resource for FAISS search instances and embedding models')
p7.add_run('\n• **Metadata Preloading:** Preload all metadata to reduce database query overhead')
p7.add_run('\n• **Deletion Handling:** Filter deleted entries from search results at query time')

doc.add_heading('3.3 BM25 Retriever Optimization', level=2)
p8 = doc.add_paragraph()
p8.paragraph_format.first_line_indent = Cm(0.74)
p8.add_run('**Problem:** nltk data download caused network timeout issues in China')
p8.add_run('\n**Solution:** Implemented built-in stopword list and simple tokenizer without external dependencies')

doc.add_heading('3.4 Search Performance Improvements', level=2)
p9 = doc.add_paragraph()
p9.paragraph_format.first_line_indent = Cm(0.74)
p9.add_run('**Pre-optimization:** Search response time exceeded 5 seconds for complex queries')
p9.add_run('\n**Post-optimization:** Pure_BM25: 0.009s, Pure_Vector: 0.027s')
p9.add_run('\n\n**Optimizations Applied:**')
p9.add_run('\n• Model instance caching')
p9.add_run('\n• Metadata batch loading')
p9.add_run('\n• Reduced database round-trips')

doc.add_heading('4. API Provider Change', level=1)
p10 = doc.add_paragraph()
p10.paragraph_format.first_line_indent = Cm(0.74)
p10.add_run('**Original API:** ')
p10.add_run('Gemini API').bold = True
p10.add_run(' for LLM-based reranking')
p10.add_run('\n**Problem:** Network connectivity instability in China resulted in frequent API call failures')
p10.add_run('\n**Solution:** Replaced with ')
p10.add_run('DeepSeek API').bold = True
p10.add_run(' (deepseek-chat model)')
p10.add_run('\n\n**Benefits:**')
p10.add_run('\n• Better network connectivity from China')
p10.add_run('\n• Lower latency for API calls')
p10.add_run('\n• Cost-effective pricing model')
p10.add_run('\n• Open-source friendly terms')

doc.save('Weekly_Report.docx')
print('✅ Weekly_Report.docx saved successfully')
