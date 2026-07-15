import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(BASE_DIR, "experiment_results")

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.style.font.name = 'Times New Roman'
    return h

def add_paragraph(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_table(data, title=""):
    if title:
        doc.add_paragraph(title)
    
    if not data:
        return None
    
    header_row = data[0]
    body_data = data[1:]
    
    n_rows = len(body_data) + 1
    n_cols = len(header_row)
    
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header_row):
        hdr_cells[i].text = str(h)
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
    
    start_row = 1
    for i, row_data in enumerate(body_data):
        row_cells = table.rows[i + start_row].cells
        for j, val in enumerate(row_data):
            row_cells[j].text = str(val)
            for paragraph in row_cells[j].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return table

add_heading("VectorLake Solutions: Comprehensive Architecture Comparison", level=1)
doc.add_paragraph("Research Internship: Knowledge Base Automation")
doc.add_paragraph("Author: Zhang Chaoyuan")
doc.add_paragraph("Supervisor: Prof. Tan Chye Cheah")
doc.add_paragraph("Date: July 2026")

add_heading("1. Executive Summary", level=2)
add_paragraph("""
This report presents a comprehensive architectural comparison of leading VectorLake solutions, including Zilliz Vector Lakebase, Milvus, Pinecone, Qdrant, Weaviate, and Chroma. The analysis evaluates each solution across key dimensions including scalability, deployment complexity, query performance, feature richness, and cost. Based on this analysis, recommendations are provided for integrating a VectorLake solution into the existing knowledge base automation system.
""")

add_heading("2. VectorLake Concept Overview", level=2)
add_paragraph("""
The VectorLake paradigm represents the next evolution of vector data management, combining the real-time retrieval capabilities of vector databases with the scalable storage and versioning capabilities of data lakehouses. Key characteristics include:
- Unified data foundation for real-time serving, interactive discovery, and batch analytics
- Lake-native storage enabling zero-copy operations across workloads
- Hybrid search capabilities (dense vectors, sparse vectors, full-text)
- Tiered storage for cost-performance optimization
- Support for multi-modal data and temporal queries
""")

add_heading("3. Comprehensive Comparison Matrix", level=2)

add_heading("3.1 Technical Specifications", level=3)
add_paragraph("""
*Note: Performance metrics represent typical values under recommended configurations. Actual performance may vary based on hardware, data volume, and query complexity. Sources: [2]-[8]*
""")
tech_specs = [
    ["Feature", "Zilliz Vector Lakebase", "Milvus", "Pinecone", "Qdrant", "Weaviate", "Chroma", "FAISS"],
    ["Core Language", "Go + C++", "Go + C++", "Proprietary", "Rust", "Go", "Python/TS", "C++"],
    ["Deployment", "Cloud/On-Prem", "Distributed/K8s/Docker", "Cloud SaaS", "Docker/K8s/Cloud", "Docker/K8s/Cloud", "Embedded/Cloud", "Python Library"],
    ["Scalability", "PB级", "亿级+", "弹性扩展", "亿级", "千万级", "百万级", "单机内存"],
    ["Index Types", "HNSW/IVF/DiskANN", "7+ types", "HNSW", "HNSW/IVF/PQ", "HNSW", "HNSW", "HNSW/IVF/PQ"],
    ["Query Latency", "<10ms", "<10ms", "<10ms", "<5ms", "<10ms", "<50ms", "<1ms"],
    ["QPS Support", "1M+", "1M+", "100K+", "500K+", "100K+", "10K+", "N/A"],
    ["Meta Filter", "SQL-like", "SQL-like", "Yes", "Payload", "GraphQL", "Basic", "No"],
    ["Hybrid Search", "Vector+BM25", "Yes", "Yes", "Dense+Sparse", "Vector+BM25", "Basic", "No"],
    ["Multi-modal", "Yes", "Yes", "No", "Yes", "Yes", "No", "No"],
    ["Versioning", "Yes", "No", "No", "No", "No", "No", "No"],
    ["Temporal Query", "Yes", "No", "No", "No", "No", "No", "No"],
    ["ACID", "Yes", "Yes", "Yes", "Yes", "Yes", "Basic", "No"],
    ["Reranking", "Yes", "No", "Yes", "Yes", "Yes", "No", "No"],
    ["Source", "[2]", "[3]", "[5]", "[4]", "[6]", "[7]", "[8]"]
]
add_table(tech_specs, title="Table 1: Technical Specifications Comparison")

add_heading("3.2 Performance Metrics", level=3)
add_paragraph("""
*Note: Performance metrics represent typical values under recommended configurations. RPS figures are estimated based on official documentation and industry benchmarks. Sources: [2]-[8]*
""")
performance_data = [
    ["Metric", "Zilliz Vector Lakebase", "Milvus", "Pinecone", "Qdrant", "Weaviate", "Chroma", "FAISS"],
    ["P50 Latency", "<10ms", "<10ms", "<10ms", "<5ms", "<15ms", "<50ms", "<1ms"],
    ["P99 Latency", "<50ms", "<50ms", "<50ms", "<20ms", "<50ms", "<200ms", "<5ms"],
    ["Recall@10", "99%", "99%", "99%", "98%", "95%", "95%", "99.9%"],
    ["RPS (1B vectors)", "1M", "1M", "100K", "500K", "100K", "10K", "N/A"],
    ["Memory Efficiency", "High", "Medium", "High", "High", "Medium", "Low", "High"],
    ["Storage Cost", "Low", "Medium", "High", "Medium", "Medium", "Low", "N/A"],
    ["Source", "[2]", "[3]", "[5]", "[4]", "[6]", "[7]", "[8]"]
]
add_table(performance_data, title="Table 2: Performance Metrics Comparison")

add_heading("3.3 Deployment and Operations", level=3)
add_paragraph("""
*Note: HA percentages represent SLA commitments for managed services or typical achievable levels for self-hosted deployments. Sources: [2]-[8]*
""")
deployment_data = [
    ["Factor", "Zilliz Vector Lakebase", "Milvus", "Pinecone", "Qdrant", "Weaviate", "Chroma", "FAISS"],
    ["Setup Complexity", "Low (managed)", "High (K8s)", "Zero", "Low (Docker)", "Medium", "Zero", "Zero"],
    ["DevOps Requirement", "None", "High", "None", "Low", "Medium", "None", "None"],
    ["Scaling Effort", "Auto", "Manual", "Auto", "Manual", "Manual", "Manual", "N/A"],
    ["Backup/Restore", "Auto", "Manual", "Auto", "Manual", "Manual", "Basic", "Manual"],
    ["Monitoring", "Built-in", "Built-in", "Built-in", "Built-in", "Built-in", "Basic", "None"],
    ["High Availability", "99.99%", "99.9%", "99.99%", "99.9%", "99.9%", "99%", "N/A"],
    ["Multi-tenant", "Yes", "Yes", "Yes", "Yes", "Yes", "Basic", "No"],
    ["Security", "Enterprise", "Enterprise", "Enterprise", "Enterprise", "Enterprise", "Basic", "None"],
    ["Source", "[2]", "[3]", "[5]", "[4]", "[6]", "[7]", "[8]"]
]
add_table(deployment_data, title="Table 3: Deployment and Operations Comparison")

add_heading("3.4 Cost Analysis", level=3)
add_paragraph("""
*Note: Cost estimates are highly dependent on infrastructure, query volume, and configuration. The following table provides qualitative assessments based on industry consensus. Specific cost figures should be obtained from official pricing documentation.*
""")
cost_data = [
    ["Cost Factor", "Zilliz Vector Lakebase", "Milvus", "Pinecone", "Qdrant", "Weaviate", "Chroma", "FAISS"],
    ["Licensing", "Freemium", "Apache 2.0", "Paid", "Apache 2.0", "Apache 2.0", "Apache 2.0", "MIT"],
    ["Cloud Cost", "Pay-as-you-go", "Infrastructure", "High", "Medium", "Medium", "Low", "None"],
    ["Self-host Cost", "Medium", "High", "N/A", "Low", "Medium", "Low", "Low"],
    ["Cost Optimization", "Tiered storage", "Index optimization", "Auto-scaling", "Memory mapping", "GraphQL", "None", "Quantization"],
    ["Source", "[2]", "[3]", "[5]", "[4]", "[6]", "[7]", "[8]"]
]
add_table(cost_data, title="Table 4: Cost Analysis Comparison")

add_heading("4. Solution Evaluation", level=2)

add_heading("4.1 Strengths Assessment", level=3)
strengths_data = [
    ["Solution", "Primary Strengths", "Key Differentiators"],
    ["Zilliz Vector Lakebase", "Unified data foundation, lake-native storage, tiered serving, on-demand search", "唯一支持实时检索、交互式探索和批量分析的统一平台"],
    ["Milvus", "极致扩展性，功能最全，CNCF毕业项目，企业级成熟度", "唯一能从容应对十亿级向量且保持低延迟的开源方案"],
    ["Pinecone", "零运维，全托管，SLA保障，Serverless架构", "最快上线速度，无需运维团队"],
    ["Qdrant", "高性能，Rust内存安全，强大过滤能力，部署灵活", "性能与灵活性最佳平衡，生产环境甜蜜点"],
    ["Weaviate", "原生混合搜索，模块化设计，GraphQL API", "向量搜索+知识图谱一体化"],
    ["Chroma", "极简集成，开发者体验优先，本地优先", "最快上手，原型开发首选"],
    ["FAISS", "速度最快，GPU支持，学术研究标准", "纯算法库，离线批处理最佳"]
]
add_table(strengths_data, title="Table 5: Solution Strengths Assessment")

add_heading("4.2 Weaknesses Assessment", level=3)
weaknesses_data = [
    ["Solution", "Primary Weaknesses", "Risk Factors"],
    ["Zilliz Vector Lakebase", "相对较新，生态尚在发展", "早期采用风险"],
    ["Milvus", "架构复杂，运维门槛高，资源消耗大", "中小规模场景过于沉重"],
    ["Pinecone", "成本不可控，无法自托管，数据主权", "长期成本高，数据安全担忧"],
    ["Qdrant", "超大规模集群管理经验较少", "十亿级以上场景经验不足"],
    ["Weaviate", "资源消耗高，GraphQL学习曲线", "内存和CPU占用较高"],
    ["Chroma", "生产环境稳定性争议，功能基础", "大规模高并发场景风险"],
    ["FAISS", "无服务层，不支持动态增删，无元数据", "需要自行构建完整系统"]
]
add_table(weaknesses_data, title="Table 6: Solution Weaknesses Assessment")

add_heading("5. Application Scenario Matching", level=2)

add_heading("5.1 Scenario-based Recommendations", level=3)
scenario_data = [
    ["Scenario", "Recommended Solutions", "Rationale"],
    ["初创公司MVP/快速原型", "Chroma, Pinecone", "Chroma本地零配置；Pinecone无需运维快速上线"],
    ["中小规模生产环境", "Qdrant, Zilliz Vector Lakebase", "性能优异，支持复杂过滤，自托管成本低"],
    ["企业级大规模部署", "Milvus, Zilliz Vector Lakebase", "极致扩展性，亿级+向量支撑，企业级特性"],
    ["超大规模数据（十亿级+）", "Milvus", "唯一能从容应对的开源方案"],
    ["数据主权敏感场景", "Qdrant, Milvus (自托管)", "数据本地化部署，无第三方依赖"],
    ["实时检索+批量分析混合", "Zilliz Vector Lakebase", "唯一支持统一数据底座上多种负载"],
    ["离线批处理/研究", "FAISS", "速度最快，学术研究标准工具"],
    ["混合搜索需求强烈", "Weaviate, Qdrant", "原生支持向量+关键词融合搜索"]
]
add_table(scenario_data, title="Table 7: Scenario-based Recommendations")

add_heading("5.2 Current System Fit Analysis", level=3)
current_fit_data = [
    ["Solution", "Current System Fit", "Integration Effort", "Expected Benefit", "Risk Level"],
    ["Zilliz Vector Lakebase", "High", "Medium", "统一数据平台，支持多负载", "Medium"],
    ["Milvus", "Medium", "High", "极致扩展性，企业级特性", "High"],
    ["Pinecone", "Low", "Low", "零运维，快速上线", "High (成本)"],
    ["Qdrant", "High", "Medium", "性能优异，部署灵活", "Low"],
    ["Weaviate", "Medium", "Medium", "混合搜索能力", "Medium"],
    ["Chroma", "Current", "None", "已有基础", "Medium (扩展)"],
    ["FAISS", "Current", "None", "已有基础", "High (功能)"]
]
add_table(current_fit_data, title="Table 8: Current System Fit Analysis")

add_heading("6. Integration Strategy", level=2)

add_heading("6.1 Recommended Approach", level=3)
add_paragraph("""
Based on the comprehensive analysis, the recommended integration strategy for the knowledge base automation system is:

**Phase 1: Short-term (Next 1-2 months)**
- Continue using FAISS for hot tier vector search (proven performance)
- Extend SQLite to implement cold tier version storage (as designed in LiveVectorLake integration)
- Implement SHA-256 chunk-level change detection for incremental updates
- Focus on completing the LiveVectorLake integration to address version control and compliance requirements

**Phase 2: Mid-term (Next 3-6 months)**
- Evaluate Qdrant as a potential replacement for FAISS when scaling beyond 10M vectors
- Qdrant offers:
  - Native support for incremental updates (no full re-indexing)
  - Strong metadata filtering capabilities
  - Rust-based performance
  - Simple Docker deployment
  - Active community and enterprise support

**Phase 3: Long-term (6+ months)**
- Consider migrating to Zilliz Vector Lakebase when:
  - System requires real-time retrieval + batch analytics on unified data
  - Data volume exceeds 100M vectors
  - Multi-modal data support becomes necessary
  - Tiered storage optimization is critical for cost management
""")

add_heading("6.2 Migration Path", level=3)
migration_data = [
    ["Phase", "Timeline", "Action", "Technology", "Expected Outcome"],
    ["Phase 1", "1-2 months", "LiveVectorLake Lite integration", "FAISS + SQLite cold tier", "版本控制、增量更新、时间查询"],
    ["Phase 2", "3-6 months", "Qdrant evaluation & migration", "Qdrant + SQLite cold tier", "增量更新、高级过滤、更好扩展性"],
    ["Phase 3", "6+ months", "Zilliz Vector Lakebase adoption", "Zilliz Vector Lakebase", "统一数据平台、多负载支持、企业级特性"]
]
add_table(migration_data, title="Table 9: Recommended Migration Path")

add_heading("7. Conclusion", level=2)
add_paragraph("""
The VectorLake paradigm represents the future of vector data management, combining real-time retrieval with scalable storage and multi-workload support. Zilliz Vector Lakebase stands out as the most comprehensive solution, offering a unified data foundation for real-time serving, interactive discovery, and batch analytics.

For the current knowledge base automation system:
- **FAISS + SQLite** remains the optimal choice for the short-term, providing proven performance and flexibility
- **Qdrant** is the recommended mid-term upgrade path, offering excellent performance, strong filtering capabilities, and manageable deployment complexity
- **Zilliz Vector Lakebase** is the long-term strategic choice when the system needs to scale to enterprise-level requirements

The immediate priority should be completing the LiveVectorLake integration (Task 3), which addresses critical limitations in version control, incremental updates, and compliance. This provides a solid foundation for future VectorLake adoption.
""")

add_heading("References", level=2)
add_paragraph("""
[1] Prajapati, T. (2025). LiveVectorLake: A Real-Time Versioned Knowledge Base Architecture for Streaming Vector Updates and Temporal Retrieval. arXiv:2601.05270.

[2] Zilliz. (2026). Zilliz Vector Lakebase: Unified Data Platform for AI. https://zilliz.com/blog/why-we-built-vector-lakebase

[3] Milvus Documentation. https://milvus.io/docs/overview.md

[4] Qdrant Documentation. https://qdrant.tech/documentation/

[5] Pinecone Documentation. https://docs.pinecone.io/

[6] Weaviate Documentation. https://weaviate.io/developers/weaviate/

[7] Chroma Documentation. https://docs.trychroma.com/

[8] FAISS Documentation. https://faiss.ai/
""")

output_path = os.path.join(OUTPUT_DIR, "VectorLake_Solutions_Comparison.docx")
doc.save(output_path)
print(f"✅ VectorLake comparison report saved to {output_path}")