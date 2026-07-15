import os
import json
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(BASE_DIR, "experiment_results")

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.style.font.name = 'Times New Roman'
    return h

def add_paragraph(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_table(data, title=""):
    if title:
        doc.add_paragraph(title)
    
    if not data:
        return None
    
    header_row = data[0]
    body_data = data[1:]
    
    n_rows = len(body_data) + 1
    n_cols = len(header_row)
    
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header_row):
        hdr_cells[i].text = str(h)
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
    
    start_row = 1
    for i, row_data in enumerate(body_data):
        row_cells = table.rows[i + start_row].cells
        for j, val in enumerate(row_data):
            row_cells[j].text = str(val)
            for paragraph in row_cells[j].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return table

add_heading("LiveVectorLake Architecture Analysis and Integration Design", level=1)
doc.add_paragraph("Research Internship: Knowledge Base Automation")
doc.add_paragraph("Author: Zhang Chaoyuan")
doc.add_paragraph("Supervisor: Prof. Tan Chye Cheah")
doc.add_paragraph("Date: July 2026")

add_heading("1. Executive Summary", level=2)
add_paragraph("""
This report presents a comprehensive analysis of the LiveVectorLake framework (arXiv:2601.05270) and proposes an integration design for the existing knowledge base automation system. LiveVectorLake introduces a dual-tier temporal knowledge base architecture that enables real-time semantic search on current knowledge while maintaining complete version history for compliance, auditability, and point-in-time retrieval. The core innovations include: (1) Content-addressable chunk-level synchronization using SHA-256 hashing for deterministic change detection; (2) Dual-tier storage separating hot-tier vector indices from cold-tier columnar versioning; (3) Temporal query routing enabling point-in-time knowledge retrieval with ACID consistency.
""")

add_heading("2. LiveVectorLake Core Architecture", level=2)

add_heading("2.1 Five-Layer Architecture", level=3)
add_paragraph("""
LiveVectorLake implements a five-layer architecture:
- **Layer 1: Change Detection and Ingestion**: Semantic chunking at paragraph boundaries, SHA-256 content-addressable hashing, and in-memory hash store for sub-millisecond CDC comparison.
- **Layer 2: Embedding Generation**: Selective processing of only modified/new chunks with temporal metadata (valid_from, valid_to, version_number).
- **Layer 3: Dual-Tier Storage**: Hot tier (Milvus with HNSW indexing) for current queries; Cold tier (Delta Lake with Parquet) for complete version history.
- **Layer 4: Query Engine**: Automatic routing based on query type, with validity filtering to prevent temporal leakage.
- **Layer 5: Interfaces**: CLI for batch operations and Streamlit UI with version timeline visualization.
""")

add_heading("2.2 Core Technical Contributions", level=3)
core_contributions = [
    ["Contribution", "Description", "Impact"],
    ["SHA-256 Chunk-level CDC", "Content-addressable hashing for deterministic change detection at paragraph granularity", "85-90% compute savings by embedding only modified content"],
    ["Dual-Tier Storage", "Hot tier (Milvus) for sub-100ms queries; Cold tier (Delta Lake) for version history", "90% storage reduction in hot tier; 100% version retention"],
    ["Temporal Query Engine", "Query classifier distinguishing current vs. temporal queries with validity filtering", "0% temporal leakage; ACID consistency across tiers"],
    ["Write-Ahead Logging", "Compensating transactions for cross-tier consistency", "Zero data loss across tier failures"]
]
add_table(core_contributions, title="Table 1: LiveVectorLake Core Technical Contributions")

add_heading("2.3 Performance Results", level=3)
add_paragraph("""
*Source: Prajapati, T. (2025). LiveVectorLake: A Real-Time Versioned Knowledge Base Architecture for Streaming Vector Updates and Temporal Retrieval. arXiv:2601.05270*
""")
performance_data = [
    ["Metric", "Value", "Significance"],
    ["Content Reprocessed", "10-15%", "vs. 100% for full re-indexing"],
    ["Current Query Latency", "Sub-100ms", "Interactive real-time response"],
    ["Historical Query Latency", "Sub-2s", "Acceptable for compliance audits"],
    ["Change Detection Accuracy", "100%", "Zero false positives via cryptographic hashing"],
    ["Hot Tier Storage Reduction", "90%", "Only active chunks in expensive vector DB"],
    ["Temporal Query Accuracy", "100%", "No temporal leakage"]
]
add_table(performance_data, title="Table 2: LiveVectorLake Performance Results")

add_heading("3. Current System Architecture Analysis", level=2)

add_heading("3.1 System Overview", level=3)
add_paragraph("""
The current knowledge base automation system employs the following technology stack:
- **Vector Index**: FAISS (Flat or IVF+PQ compressed index)
- **Data Storage**: SQLite/Chroma database with 2396 embeddings and 30611 metadata entries
- **Embedding Model**: BAAI/bge-small-en-v1.5 (384-dimensional vectors)
- **Retrieval**: Pure vector search, BM25 search, and hybrid RRF search
- **Reranking**: CrossEncoder (ms-marco-MiniLM-L-6-v2 and L-12-v2) and Gemini API
- **Document Processing**: Unified extractor supporting PDF, DOCX, XLSX, PPTX, images, and text files
""")

add_heading("3.2 Architecture Comparison", level=3)
comparison_data = [
    ["Feature", "LiveVectorLake", "Current System", "Gap"],
    ["Change Detection", "SHA-256 chunk-level CDC", "None (full re-indexing)", "Critical: 85% compute waste"],
    ["Version Control", "Complete temporal history", "None (overwrites only)", "Critical: No compliance support"],
    ["Dual-Tier Storage", "Hot (Milvus) + Cold (Delta Lake)", "Single-tier (FAISS + SQLite)", "Medium: No storage optimization"],
    ["Temporal Query", "Point-in-time retrieval", "Current state only", "Critical: No audit capability"],
    ["ACID Consistency", "Write-ahead logging", "Basic SQLite transactions", "Medium: No cross-tier consistency"],
    ["Incremental Update", "Selective embedding", "Full re-embedding", "Critical: High latency updates"],
    ["Query Routing", "Automatic tier selection", "Single index query", "Medium: No query optimization"]
]
add_table(comparison_data, title="Table 3: Architecture Comparison - LiveVectorLake vs. Current System")

add_heading("3.3 Current System Limitations", level=3)
limitations_data = [
    ["Limitation", "Impact", "Severity"],
    ["No Version History", "Cannot answer temporal queries; compliance violations", "Critical"],
    ["Full Re-Indexing", "85-95% wasted compute on unchanged content", "Critical"],
    ["No Change Detection", "Cannot identify which content modified", "High"],
    ["No Audit Trail", "Cannot reconstruct knowledge state at specific time", "High"],
    ["Single-Tier Storage", "No cost optimization for historical data", "Medium"],
    ["No Temporal Leakage Protection", "Historical queries may return future content", "Medium"]
]
add_table(limitations_data, title="Table 4: Current System Limitations")

add_heading("4. Integration Design: Simplified LiveVectorLake", level=2)

add_heading("4.1 Design Principles", level=3)
add_paragraph("""
The integration design follows these principles:
1. **Minimal Disruption**: Use existing FAISS+SQLite tech stack; no external dependencies
2. **Gradual Adoption**: Implement core features incrementally
3. **Cost-Effective**: No need for distributed systems (Milvus/Delta Lake) in current scale
4. **Backward Compatible**: Existing queries continue to work without modification
5. **Extensible**: Design allows future migration to full LiveVectorLake architecture
""")

add_heading("4.2 Proposed Architecture", level=3)
add_paragraph("""
The proposed simplified LiveVectorLake architecture maintains the existing FAISS index as the hot tier and extends SQLite to serve as the cold tier for version history:

**Hot Tier (FAISS)**: 
- Stores only active (current) chunks
- Optimized for low-latency semantic search
- Uses IVF+PQ compression for memory efficiency

**Cold Tier (SQLite Extension)**:
- Stores complete version history with temporal metadata
- Supports point-in-time retrieval via validity filtering
- Append-only writes for audit trail integrity

**Change Detection Layer**:
- SHA-256 hashing for chunk-level content fingerprinting
- In-memory hash store for sub-millisecond change comparison
- Selective embedding of only modified chunks

**Query Engine**:
- Automatic routing: current queries to hot tier, temporal queries to cold tier
- Validity filtering to prevent temporal leakage
""")

add_heading("4.3 Database Schema Extension", level=3)
schema_data = [
    ["Table", "Column", "Type", "Description"],
    ["chunk_versions", "chunk_id", "TEXT (SHA-256)", "Content-addressable chunk identifier"],
    ["chunk_versions", "doc_id", "INTEGER", "Reference to original document"],
    ["chunk_versions", "position", "INTEGER", "Paragraph index in document"],
    ["chunk_versions", "content", "TEXT", "Chunk content"],
    ["chunk_versions", "embedding", "BLOB", "Vector embedding (384-dim)"],
    ["chunk_versions", "valid_from", "TIMESTAMP", "Version activation time"],
    ["chunk_versions", "valid_to", "TIMESTAMP", "Version superseded time (NULL=current)"],
    ["chunk_versions", "version_number", "INTEGER", "Monotonic sequence number"],
    ["chunk_versions", "status", "TEXT", "'active'/'superseded'/'deleted'"],
    ["chunk_versions", "change_type", "TEXT", "'insert'/'update'/'delete'"],
    ["doc_hash_store", "doc_id", "INTEGER", "Document identifier"],
    ["doc_hash_store", "chunk_hashes", "JSON", "Array of SHA-256 hashes for each chunk"],
    ["doc_hash_store", "last_updated", "TIMESTAMP", "Last ingestion time"]
]
add_table(schema_data, title="Table 5: Proposed SQLite Schema Extension for Cold Tier")

add_heading("4.4 Change Detection Algorithm", level=3)
add_paragraph("""
**Algorithm: Chunk-Level Change Detection**

1. **Semantic Chunking**: Split document at paragraph boundaries (double newlines)
2. **Normalization**: Apply whitespace stripping and case-folding to each chunk
3. **SHA-256 Hashing**: Compute content-addressable hash for each chunk:
   ```
   chunk_id = SHA256(normalize(content))
   ```
4. **Hash Comparison**: Compare new hashes against stored hashes from previous version
5. **Classification**: 
   - New: Hash not in previous version
   - Modified: Different hash at same position
   - Deleted: Hash absent in new version
   - Unchanged: Hash present, same position
6. **Selective Embedding**: Only embed new and modified chunks
7. **Version Update**: Mark old versions as 'superseded', insert new versions
""")

add_heading("4.5 Query Processing Flow", level=3)
query_flow_data = [
    ["Step", "Current Query", "Temporal Query"],
    ["1. Query Classification", "Detect no temporal constraint", "Detect timestamp in query"],
    ["2. Routing", "Send to FAISS hot tier", "Send to SQLite cold tier"],
    ["3. Search", "HNSW approximate nearest neighbor", "Vector search with validity filtering"],
    ["4. Validity Filter", "Not applicable", "Filter chunks where valid_from <= timestamp < valid_to"],
    ["5. Result Return", "Current active content", "Historical content at requested time"],
    ["6. Temporal Leakage", "Not applicable", "Guaranteed zero leakage"]
]
add_table(query_flow_data, title="Table 6: Query Processing Flow Comparison")

add_heading("4.6 Expected Improvements", level=3)
improvements_data = [
    ["Metric", "Before (Current System)", "After (LiveVectorLake Integration)", "Improvement"],
    ["Content Reprocessed", "100% (full re-index)", "10-15% (chunk-level CDC)", "85-90% reduction"],
    ["Update Latency", "Minutes (full embedding)", "Seconds (selective embedding)", "90% reduction"],
    ["Version History", "None", "Complete (all versions)", "New capability"],
    ["Temporal Query", "Not supported", "Supported (point-in-time)", "New capability"],
    ["Compliance Audit", "Not possible", "Full audit trail", "New capability"],
    ["Storage Efficiency", "All data in hot tier", "90% in cold tier", "Cost optimization"],
    ["Change Detection", "Not available", "100% accurate (SHA-256)", "New capability"]
]
add_table(improvements_data, title="Table 7: Expected Improvements from LiveVectorLake Integration")

add_heading("5. Implementation Roadmap", level=2)

add_heading("5.1 Phase 1: Core Infrastructure (Weeks 1-2)", level=3)
add_paragraph("""
- Implement SHA-256 chunk hashing module
- Create SQLite cold tier tables (chunk_versions, doc_hash_store)
- Develop change detection logic for document ingestion
- Integrate selective embedding into existing pipeline
""")

add_heading("5.2 Phase 2: Version Management (Weeks 3-4)", level=3)
add_paragraph("""
- Implement version tracking with valid_from/valid_to timestamps
- Develop append-only write operations for cold tier
- Create version history visualization in UI
- Implement basic audit trail generation
""")

add_heading("5.3 Phase 3: Query Engine (Weeks 5-6)", level=3)
add_paragraph("""
- Develop query classifier for current vs. temporal queries
- Implement temporal query routing to cold tier
- Add validity filtering to prevent temporal leakage
- Optimize cold tier vector search performance
""")

add_heading("5.4 Phase 4: Evaluation & Optimization (Weeks 7-8)", level=3)
add_paragraph("""
- Conduct performance evaluation (update efficiency, query latency)
- Optimize cold tier storage (compression, indexing)
- Test temporal query accuracy (0% leakage verification)
- Generate comprehensive evaluation report
""")

add_heading("6. Risk Assessment", level=2)
risk_data = [
    ["Risk", "Likelihood", "Impact", "Mitigation Strategy"],
    ["SQLite performance for cold tier queries", "Medium", "High", "Implement indexing; migrate to Delta Lake at scale"],
    ["Memory usage for hash store", "Low", "Medium", "Persistent hash store to disk; periodic cleanup"],
    ["Schema migration complexity", "Medium", "Medium", "Backward-compatible schema extension; incremental rollout"],
    ["Temporal query latency", "Medium", "Medium", "Acceptable for audit use cases; optimize with indexing"],
    ["Embedding consistency across versions", "Low", "High", "Version-lock embedding model; store model metadata"]
]
add_table(risk_data, title="Table 8: Risk Assessment")

add_heading("7. Conclusion", level=2)
add_paragraph("""
LiveVectorLake represents a significant advancement in RAG system architecture, addressing critical limitations of current systems through its dual-tier storage, chunk-level change detection, and temporal query support. The proposed integration design adapts LiveVectorLake's core concepts to the existing FAISS+SQLite tech stack, enabling gradual adoption without requiring distributed systems like Milvus or Delta Lake.

The expected outcomes include:
- 85-90% reduction in embedding compute through selective processing
- Complete version history for compliance and audit
- Point-in-time retrieval capability for temporal queries
- Zero temporal leakage through validity filtering
- Cost optimization through hot/cold tier separation

This integration will position the knowledge base automation system as a production-ready solution capable of handling continuous knowledge updates while maintaining compliance requirements—essential capabilities for enterprise RAG deployments.
""")

add_heading("References", level=2)
add_paragraph("""
[1] Prajapati, T. (2025). LiveVectorLake: A Real-Time Versioned Knowledge Base Architecture for Streaming Vector Updates and Temporal Retrieval. arXiv:2601.05270.

[2] Lewis, P., et al. (2020). Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. NeurIPS.

[3] Zilliz. (2026). Zilliz Vector Lakebase: Unified Data Platform for AI. https://zilliz.com/blog/why-we-built-vector-lakebase

[4] Milvus Documentation. https://milvus.io/docs/overview.md

[5] Delta Lake Documentation. https://docs.delta.io/latest/index.html
""")

output_path = os.path.join(OUTPUT_DIR, "LiveVectorLake_Analysis_and_Integration_Design.docx")
doc.save(output_path)
print(f"✅ LiveVectorLake analysis report saved to {output_path}")